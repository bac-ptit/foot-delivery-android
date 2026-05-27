#!/usr/bin/env python3
"""
Xuất báo cáo kỹ thuật Foot Delivery Android ra file .docx
Font body: 13pt, Header: 14pt, nền trắng, có đánh số trang, không header trang.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Cấu hình ──────────────────────────────────────────────────

BODY_FONT_SIZE = Pt(13)
HEADER_FONT_SIZE = Pt(14)
TITLE_FONT_SIZE = Pt(18)
SUBTITLE_FONT_SIZE = Pt(15)
FONT_NAME = "Times New Roman"
FONT_NAME_MONO = "Consolas"

# ── Hàm tiện ích ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Đặt màu nền cho ô bảng."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Đặt border cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_page_number(section):
    """Thêm đánh số trang vào footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run("1")
    run4.font.size = Pt(10)

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    # Thêm " / " và tổng số trang
    run6 = paragraph.add_run(" / ")
    run6.font.size = Pt(10)

    run7 = paragraph.add_run()
    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run7._r.append(fldChar4)

    run8 = paragraph.add_run()
    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run8._r.append(instrText2)

    run9 = paragraph.add_run()
    fldChar5 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run9._r.append(fldChar5)

    run10 = paragraph.add_run("1")
    run10.font.size = Pt(10)

    run11 = paragraph.add_run()
    fldChar6 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run11._r.append(fldChar6)


def add_paragraph(doc, text, style=None, bold=False, font_size=None, alignment=None, space_after=None, space_before=None, font_name=None, color=None):
    """Thêm paragraph với style tùy chỉnh."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = font_size or BODY_FONT_SIZE
    run.font.name = font_name or FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or FONT_NAME)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    return p


def add_heading_styled(doc, text, level=1):
    """Thêm heading với style đã định nghĩa."""
    p = doc.add_paragraph()
    run = p.add_run(text)

    if level == 0:
        run.font.size = TITLE_FONT_SIZE
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
    elif level == 1:
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        run.font.size = SUBTITLE_FONT_SIZE
        run.bold = True
        run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        run.font.size = HEADER_FONT_SIZE
        run.bold = True
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Thêm bảng với header màu xanh, các dòng xen kẽ trắng/xám."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = BODY_FONT_SIZE
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = BODY_FONT_SIZE
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

            # Xen kẽ màu
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F0F4FF")
            else:
                set_cell_shading(cell, "FFFFFF")

    # Đặt chiều rộng cột nếu có
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_code_block(doc, code_text):
    """Thêm block code với nền xám."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Tạo border nền xám
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)

    # Thêm indent
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)

    run = p.add_run(code_text)
    run.font.name = FONT_NAME_MONO
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_bullet(doc, text, level=0):
    """Thêm bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = BODY_FONT_SIZE
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


def add_flow_arrow(doc, items):
    """Thêm flowchart dạng text với mũi tên."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        run = p.add_run(item)
        run.font.name = FONT_NAME_MONO
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)


# ══════════════════════════════════════════════════════════════
# TẠO DOCUMENT
# ══════════════════════════════════════════════════════════════

def create_report():
    doc = Document()

    # ── Cấu hình trang ────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Xóa header mặc định
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""

    # Thêm đánh số trang
    add_page_number(section)

    # ── Trang bìa ─────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    add_paragraph(doc, "BÁO CÁO KỸ THUẬT", bold=True, font_size=Pt(28),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1A, 0x56, 0xDB))

    add_paragraph(doc, "FOOT DELIVERY ANDROID", bold=True, font_size=Pt(22),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x2D, 0x3A, 0x4A))

    doc.add_paragraph()

    add_paragraph(doc, "Ứng dụng đặt đồ ăn giao hàng", font_size=Pt(14),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x4A, 0x55, 0x68))

    doc.add_paragraph()

    # Thông tin dự án
    info_items = [
        ("Dự án", "Foot Delivery Android"),
        ("Package", "com.example.myapp"),
        ("Công nghệ", "Kotlin, Retrofit, Firebase, FastAPI"),
        ("Ngày", "27/05/2026"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(f"{label}: ")
        run1.bold = True
        run1.font.size = BODY_FONT_SIZE
        run1.font.name = FONT_NAME
        run2 = p.add_run(value)
        run2.font.size = BODY_FONT_SIZE
        run2.font.name = FONT_NAME

    # Ngắt trang
    doc.add_page_break()

    # ── Mục lục ───────────────────────────────────────────────
    add_heading_styled(doc, "MỤC LỤC", level=1)

    toc_items = [
        "PHẦN 1: TÀI LIỆU KỸ THUẬT",
        "  1. Danh sách chức năng được phân công",
        "  2. Kiến trúc chi tiết hệ thống liên quan",
        "  3. Code đáp ứng chức năng",
        "  4. Hướng dẫn cài đặt và triển khai",
        "PHẦN 2: CODE",
        "  1. Phần cá nhân thực hiện",
        "  2. Tối ưu source code",
        "  3. Kiểm tra comment",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = BODY_FONT_SIZE
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        if not item.startswith("  "):
            run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 1: TÀI LIỆU KỸ THUẬT
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, "PHẦN 1: TÀI LIỆU KỸ THUẬT", level=0)

    # ── 1. Danh sách chức năng ────────────────────────────────
    add_heading_styled(doc, "1. Danh sách chức năng được phân công", level=1)

    add_table(doc,
        headers=["STT", "Chức năng", "Mô tả", "Trạng thái"],
        rows=[
            ["1", "Đặt món từ nhà hàng địa phương",
             "Duyệt nhà hàng, xem menu, chọn món, thêm vào giỏ hàng, đặt hàng với thanh toán COD/VNPay", "Hoàn thành"],
            ["2", "Theo dõi đơn hàng",
             "Xem trạng thái đơn hàng theo thời gian thực (chờ xác nhận → đang giao → hoàn thành), lịch sử đơn hàng", "Hoàn thành"],
            ["3", "Tích điểm khách hàng",
             "Tính điểm tích lũy dựa trên tổng chi tiêu, hiển thị trên profile và trang theo dõi đơn", "Hoàn thành"],
            ["4", "Đánh giá và phản hồi",
             "Đánh giá đơn hàng已完成 (1-5 sao + bình luận), hiển thị đánh giá trên trang chi tiết món ăn", "Hoàn thành"],
        ],
        col_widths=[1.5, 4, 8, 2.5]
    )

    # ── 2. Kiến trúc chi tiết ─────────────────────────────────
    add_heading_styled(doc, "2. Kiến trúc chi tiết hệ thống liên quan", level=1)

    add_heading_styled(doc, "Kiến trúc tổng quan", level=2)

    add_code_block(doc,
        "┌──────────────────────────────────────────────────────────────┐\n"
        "│                    ANDROID APP (Kotlin)                        │\n"
        "│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐       │\n"
        "│  │   screens/   │  │  adapters/   │  │    api/      │       │\n"
        "│  │  (Activity)  │  │ (RecyclerView│  │ (Retrofit)   │       │\n"
        "│  │  30 files    │  │   Adapter)   │  │  4 files     │       │\n"
        "│  └──────┬───────┘  └──────┬───────┘  └──────┬───────┘       │\n"
        "│         └─────────────────┴──────────────────┘                │\n"
        "│                    SharedPreferences                           │\n"
        "│                    (JWT Token, User Info)                      │\n"
        "└───────────────────────────┬──────────────────────────────────┘\n"
        "                            │ HTTP/REST\n"
        "                            ▼\n"
        "┌──────────────────────────────────────────────────────────────┐\n"
        "│                  BACKEND (FastAPI + Python)                    │\n"
        "│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐       │\n"
        "│  │  main.py     │  │  models.py   │  │  schemas.py  │       │\n"
        "│  │  (Routes)    │  │ (SQLAlchemy) │  │ (Pydantic)   │       │\n"
        "│  └──────────────┘  └──────────────┘  └──────────────┘       │\n"
        "│                    PostgreSQL                                  │\n"
        "└──────────────────────────────────────────────────────────────┘"
    )

    # ── 2.1 Thành phần hệ thống ───────────────────────────────
    add_heading_styled(doc, "2.1 Thành phần hệ thống", level=2)
    add_heading_styled(doc, "Android App", level=3)

    add_table(doc,
        headers=["Thành phần", "Vai trò", "Input", "Output", "Liên quan"],
        rows=[
            ["screens/home.kt", "Màn hình chính", "GET /menu-items/", "RecyclerView món", "MenuItemAdapter, food_detail"],
            ["screens/list_restaurant.kt", "Danh sách nhà hàng", "GET /restaurants/", "RecyclerView nhà hàng", "RestaurantAdapter, restaurant_profile"],
            ["screens/restaurant_profile.kt", "Chi tiết nhà hàng", "GET /restaurants/{id}/", "Thông tin + menu", "MenuItemAdapterSmall, food_detail"],
            ["screens/food_detail.kt", "Chi tiết món ăn", "GET /menu-items/{id}/", "Món + đánh giá", "cart, pre_order"],
            ["screens/cart.kt", "Giỏ hàng", "CartItem (static)", "Tổng tiền", "order"],
            ["screens/order.kt", "Xác nhận đặt hàng", "Cart, Address, Voucher", "POST /orders/", "payment_methods, discouts"],
            ["screens/OrderTracking.kt", "Theo dõi đơn", "GET /users/{id}/orders/", "Đơn pending/completed", "OrderTrackingDetail"],
            ["screens/OrderTrackingDetail.kt", "Chi tiết đơn + đánh giá", "GET /orders/{id}/detail", "Trạng thái, form đánh giá", "Reviews"],
            ["screens/PointsDetail.kt", "Tích điểm", "GET /users/{id}/orders/", "Điểm theo đơn", "OrderTrackingDetail"],
            ["screens/profile.kt", "Hồ sơ", "GET /users/{id}/profile-summary", "Tên, điểm", "PointsDetail"],
        ],
        col_widths=[3.5, 2.5, 3.5, 3, 3.5]
    )

    add_heading_styled(doc, "Backend", level=3)

    add_table(doc,
        headers=["Thành phần", "Vai trò", "Input", "Output", "Liên quan"],
        rows=[
            ["main.py", "API Routes", "HTTP Request", "HTTP Response", "models, schemas"],
            ["models.py", "ORM Models", "—", "SQLAlchemy objects", "Database"],
            ["schemas.py", "Pydantic schemas", "Request body", "Validated data", "main.py"],
            ["notification_service.py", "Push notification", "user_id, title, body", "FCM notification", "Firebase"],
            ["auth.py", "JWT authentication", "username/password", "access_token", "main.py"],
        ],
        col_widths=[3.5, 2.5, 3.5, 3, 3.5]
    )

    # ── 2.2 Luồng xử lý ──────────────────────────────────────
    add_heading_styled(doc, "2.2 Luồng xử lý", level=2)

    # Đặt món
    add_heading_styled(doc, "Đặt món từ nhà hàng địa phương", level=3)

    add_flow_arrow(doc, [
        "Khách hàng",
        "    │",
        "    ▼",
        "[Màn hình chính] ──GET /menu-items/──→ [Hiển thị danh sách món]",
        "    │",
        "    ├── Tìm kiếm ──GET /menu-items/search/──→ [Kết quả]",
        "    ├── Chọn nhà hàng ──GET /restaurants/──→ [Danh sách]",
        "    │       ▼",
        "    │   [restaurant_profile] ──GET /restaurants/{id}/──→ [Menu]",
        "    │       ▼",
        "    │   [food_detail] ──GET /menu-items/{id}/──→ [Chi tiết]",
        "    │       ▼",
        "    │   [Thêm vào giỏ] ──→ cart.cartList (in-memory)",
        "    ▼",
        "[Giỏ hàng] ──→ Chọn món → Điều chỉnh số lượng",
        "    ▼",
        "[Xác nhận đơn] ──→ Địa chỉ + Voucher + Thanh toán",
        "    ├── COD ──→ POST /orders/ ──→ confirmed",
        "    └── VNPay ──→ GET /create-payment ──→ WebView",
        "                    ▼",
        "            vnp_ResponseCode=00? → paid / cancelled",
        "                    ▼",
        "            [payment_successful]",
    ])

    # Theo dõi đơn hàng
    add_heading_styled(doc, "Theo dõi đơn hàng", level=3)

    add_flow_arrow(doc, [
        "Khách hàng",
        "    │",
        "    ▼",
        "[OrderTracking] ──GET /users/{id}/orders/──→ [Tất cả đơn]",
        "    │",
        "    ├── Pending (đỏ): \"Chưa giao\"",
        "    │       ▼",
        "    │   [OrderTrackingDetail] ── Chi tiết + Nút đánh giá DISABLED",
        "    │",
        "    └── Completed (xanh): \"Đã giao\"",
        "            ▼",
        "        [OrderTrackingDetail] ── Chi tiết + Nút đánh giá ENABLED",
        "",
        "Backend push notification:",
        "    PUT /orders/{id}/status → notify_user()",
        "    → \"Đơn hàng được đặt thành công!\" (paid/confirmed)",
        "    → \"Đơn hàng đang giao\" (delivering)",
        "    → \"Đơn hàng đã giao\" (completed)",
    ])

    # Tích điểm
    add_heading_styled(doc, "Tích điểm khách hàng", level=3)

    add_flow_arrow(doc, [
        "Đơn hàng hoàn thành",
        "    │",
        "    ▼",
        "[profile] ──GET /users/{id}/profile-summary──→",
        "    │",
        "    │   Backend: points = total_spent // 10000",
        "    ▼",
        "[Hiển thị điểm trên profile]",
        "    ├── Click điểm → [OrderTracking] → loadPointsSummary()",
        "    └── [PointsDetail] ── Mỗi đơn: điểm = totalprice / 1000",
    ])

    # Đánh giá
    add_heading_styled(doc, "Đánh giá và phản hồi", level=3)

    add_flow_arrow(doc, [
        "[OrderTrackingDetail] (chỉ khi đơn Đã giao)",
        "    │",
        "    ├── Chọn sao (1-5) → updateStarColors()",
        "    ├── Nhập bình luận → etFeedback",
        "    └── Nhấn \"Đánh giá\"",
        "            │",
        "            ▼",
        "        Validation: rating > 0, feedback không trống",
        "            ▼",
        "        POST /reviews/ { rating, comment, orderid, userid }",
        "            ▼",
        "        Backend: Lưu Review → Tính lại avg_rating → Cập nhật restaurant.rating",
        "            ▼",
        "        Toast: \"Cảm ơn bạn đã đánh giá!\"",
    ])

    # ── 3. Code đáp ứng chức năng ─────────────────────────────
    doc.add_page_break()
    add_heading_styled(doc, "3. Code đáp ứng chức năng", level=1)

    # 3.1 File liên quan
    add_heading_styled(doc, "3.1 File liên quan", level=2)

    add_table(doc,
        headers=["File", "Vai trò", "Chức năng"],
        rows=[
            ["screens/home.kt", "Màn hình chính", "Hiển thị tất cả món ăn, tìm kiếm"],
            ["screens/list_restaurant.kt", "Danh sách nhà hàng", "Tìm kiếm và hiển thị nhà hàng"],
            ["screens/restaurant_profile.kt", "Chi tiết nhà hàng", "Thông tin nhà hàng, menu"],
            ["screens/food_detail.kt", "Chi tiết món ăn", "Thông tin món, đánh giá, thêm giỏ"],
            ["screens/cart.kt", "Giỏ hàng", "Quản lý món đã chọn"],
            ["screens/order.kt", "Xác nhận đơn", "Địa chỉ, voucher, thanh toán"],
            ["screens/OrderTracking.kt", "Theo dõi đơn", "Danh sách đơn, phân trang"],
            ["screens/OrderTrackingDetail.kt", "Chi tiết đơn + đánh giá", "Trạng thái, form đánh giá"],
            ["screens/PointsDetail.kt", "Tích điểm", "Điểm theo đơn hàng"],
            ["screens/profile.kt", "Hồ sơ", "Tên, điểm tích lũy"],
            ["adapters/MenuItemAdapter.kt", "Adapter món ăn", "Hiển thị trên home"],
            ["adapters/RestaurantAdapter.kt", "Adapter nhà hàng", "Hiển thị danh sách"],
            ["api/ApiService.kt", "API definitions", "Data classes + Retrofit interface"],
            ["api/RetrofitClient.kt", "HTTP Client", "Singleton + auth interceptor"],
        ],
        col_widths=[4, 3.5, 5]
    )

    # 3.2 Class liên quan
    add_heading_styled(doc, "3.2 Class liên quan", level=2)

    add_table(doc,
        headers=["Class", "File", "Vai trò"],
        rows=[
            ["home", "screens/home.kt", "Activity màn hình chính"],
            ["list_restaurant", "screens/list_restaurant.kt", "Activity danh sách nhà hàng"],
            ["restaurant_profile", "screens/restaurant_profile.kt", "Activity chi tiết nhà hàng"],
            ["food_detail", "screens/food_detail.kt", "Activity chi tiết món ăn"],
            ["cart", "screens/cart.kt", "Activity giỏ hàng + companion object"],
            ["CartItem", "screens/cart.kt", "Data class sản phẩm trong giỏ"],
            ["order", "screens/order.kt", "Activity xác nhận đơn"],
            ["OrderTracking", "screens/OrderTracking.kt", "Activity theo dõi đơn"],
            ["OrderTrackingDetail", "screens/OrderTrackingDetail.kt", "Activity chi tiết đơn + đánh giá"],
            ["PointsDetail", "screens/PointsDetail.kt", "Activity tích điểm"],
            ["profile", "screens/profile.kt", "Activity hồ sơ"],
            ["MenuItemAdapter", "adapters/MenuItemAdapter.kt", "Adapter món ăn"],
            ["RestaurantAdapter", "adapters/RestaurantAdapter.kt", "Adapter nhà hàng"],
            ["RetrofitClient", "api/RetrofitClient.kt", "Singleton HTTP client"],
            ["ApiService", "api/ApiService.kt", "Retrofit interface"],
        ],
        col_widths=[3, 5, 5]
    )

    # 3.3 Hàm liên quan
    add_heading_styled(doc, "3.3 Hàm liên quan", level=2)

    add_heading_styled(doc, "Đặt món", level=3)
    add_table(doc,
        headers=["Hàm", "File", "Input", "Output", "Chức năng"],
        rows=[
            ["fetchMenuItems()", "home.kt", "—", "List<MenuItem>", "Tải tất cả món ăn"],
            ["searchMenuItems(query)", "home.kt", "String", "List<MenuItem>", "Tìm kiếm món theo tên"],
            ["fetchRestaurants()", "list_restaurant.kt", "—", "List<Restaurant>", "Tải danh sách nhà hàng"],
            ["fetchRestaurantDetails(id)", "restaurant_profile.kt", "Int", "Restaurant", "Tải chi tiết nhà hàng"],
            ["fetchFoodDetails(id)", "food_detail.kt", "Int", "MenuItem", "Tải chi tiết món ăn"],
            ["addToCart()", "food_detail.kt", "MenuItem, Int", "—", "Thêm vào giỏ hàng"],
            ["createOrderFromCart()", "order.kt", "—", "OrderResponse", "Tạo đơn hàng"],
        ],
        col_widths=[3.5, 3, 2, 2.5, 3]
    )

    add_heading_styled(doc, "Theo dõi đơn hàng", level=3)
    add_table(doc,
        headers=["Hàm", "File", "Input", "Output", "Chức năng"],
        rows=[
            ["loadAllOrders()", "OrderTracking.kt", "—", "List<OrderResponse>", "Tải tất cả đơn"],
            ["loadOrdersDetail()", "OrderTracking.kt", "List, List", "—", "Tải chi tiết từng đơn"],
            ["isCompletedStatus()", "OrderTracking.kt", "String", "Boolean", "Kiểm tra trạng thái"],
            ["loadOrderSummary()", "OrderTrackingDetail.kt", "Int", "OrderDetailResponse", "Tải chi tiết đơn"],
            ["formatOrderDate()", "OrderTrackingDetail.kt", "—", "String", "Định dạng ngày"],
        ],
        col_widths=[3.5, 3, 2, 2.5, 3]
    )

    add_heading_styled(doc, "Tích điểm", level=3)
    add_table(doc,
        headers=["Hàm", "File", "Input", "Output", "Chức năng"],
        rows=[
            ["loadPointsSummary()", "OrderTracking.kt", "—", "UserProfileSummary", "Tải tổng điểm"],
            ["loadOrders()", "PointsDetail.kt", "—", "List<OrderResponse>", "Tải đơn已完成"],
            ["bindOrderCard()", "PointsDetail.kt", "OrderResponse", "—", "Hiển thị card điểm"],
        ],
        col_widths=[3.5, 3, 2, 2.5, 3]
    )

    add_heading_styled(doc, "Đánh giá", level=3)
    add_table(doc,
        headers=["Hàm", "File", "Input", "Output", "Chức năng"],
        rows=[
            ["updateStarColors()", "OrderTrackingDetail.kt", "Array<ImageView>", "—", "Cập nhật màu sao"],
            ["createReview()", "OrderTrackingDetail.kt", "ReviewCreateRequest", "ReviewResponse", "Gửi đánh giá"],
            ["displayReviews()", "food_detail.kt", "—", "—", "Hiển thị đánh giá"],
        ],
        col_widths=[3.5, 3, 2, 2.5, 3]
    )

    # 3.4 Database
    add_heading_styled(doc, "3.4 Database", level=2)

    add_table(doc,
        headers=["Bảng", "Chức năng", "Quan hệ"],
        rows=[
            ["Restaurant", "Thông tin nhà hàng", "1-N: MenuItem, Orders, Review"],
            ["MenuItem", "Món ăn trong menu", "N-1: Restaurant; 1-N: OrderItem, Review"],
            ["Orders", "Đơn hàng", "N-1: User, Restaurant, Address; 1-N: OrderItem"],
            ["OrderItem", "Chi tiết đơn hàng", "N-1: Orders, MenuItem"],
            ["User", "Người dùng", "1-N: Orders, Address, Review"],
            ["Review", "Đánh giá", "N-1: User, MenuItem, Restaurant, Orders"],
            ["Address", "Địa chỉ giao hàng", "N-1: User; 1-N: Orders"],
            ["Payment", "Thanh toán", "N-1: Orders"],
            ["Promotion", "Mã giảm giá", "1-N: UsedPromotion"],
            ["Notification", "Thông báo", "N-1: User, Orders"],
            ["Delivery", "Giao hàng", "N-1: Orders, Shipper"],
            ["LoyaltyPoint", "Điểm tích lũy (KHÔNG DÙNG)", "N-1: User, MenuItem"],
        ],
        col_widths=[3, 5, 5]
    )

    add_heading_styled(doc, "Chi tiết bảng Orders", level=3)
    add_bullet(doc, "Primary Key: id (Integer, auto-increment)")
    add_bullet(doc, "Foreign Key: restaurantid → Restaurant.id")
    add_bullet(doc, "Foreign Key: addressid → Address.id")
    add_bullet(doc, "Foreign Key: userid → User.id")
    add_bullet(doc, "Các trường: status, createdat, preorderdate, preordertime, totalprice")

    add_heading_styled(doc, "Chi tiết bảng Review", level=3)
    add_bullet(doc, "Primary Key: id (Integer, auto-increment)")
    add_bullet(doc, "Foreign Key: menuitemid → MenuItem.id")
    add_bullet(doc, "Foreign Key: restaurantid → Restaurant.id")
    add_bullet(doc, "Foreign Key: userid → User.id")
    add_bullet(doc, "Foreign Key: orderid → Orders.id")
    add_bullet(doc, "Các trường: rating (Integer), comment (Text)")

    # 3.5 API
    add_heading_styled(doc, "3.5 API", level=2)

    add_table(doc,
        headers=["Endpoint", "Method", "File xử lý", "Chức năng"],
        rows=[
            ["/restaurants/", "GET", "main.py", "Danh sách nhà hàng"],
            ["/restaurants/search/", "GET", "main.py", "Tìm kiếm nhà hàng"],
            ["/restaurants/{id}/", "GET", "main.py", "Chi tiết nhà hàng + menu"],
            ["/menu-items/", "GET", "main.py", "Danh sách tất cả món ăn"],
            ["/menu-items/{id}/", "GET", "main.py", "Chi tiết món ăn + đánh giá"],
            ["/menu-items/search/", "GET", "main.py", "Tìm kiếm món theo tên"],
            ["/orders/", "POST", "main.py", "Tạo đơn hàng mới"],
            ["/users/{id}/orders/", "GET", "main.py", "Danh sách đơn hàng"],
            ["/orders/{id}/detail", "GET", "main.py", "Chi tiết đơn hàng"],
            ["/orders/{id}/status", "PUT", "main.py", "Cập nhật trạng thái"],
            ["/reviews/", "POST", "main.py", "Tạo đánh giá mới"],
            ["/users/{id}/profile-summary", "GET", "main.py", "Hồ sơ + điểm tích lũy"],
            ["/promotions/", "GET", "main.py", "Danh sách mã giảm giá"],
            ["/create-payment", "GET", "main.py", "Tạo URL VNPay"],
            ["/users/{id}/addresses/", "GET", "main.py", "Danh sách địa chỉ"],
        ],
        col_widths=[4, 1.5, 2, 5]
    )

    # ── 4. Hướng dẫn cài đặt ──────────────────────────────────
    doc.add_page_break()
    add_heading_styled(doc, "4. Hướng dẫn cài đặt và triển khai", level=1)

    add_heading_styled(doc, "Yêu cầu môi trường", level=2)
    add_bullet(doc, "Android: Android Studio, JDK 17, Android SDK 24-34")
    add_bullet(doc, "Backend: Python 3.10+, pip/uv")
    add_bullet(doc, "Database: PostgreSQL")
    add_bullet(doc, "Docker (tùy chọn): Docker Desktop")
    add_bullet(doc, "Firebase: Tài khoản Firebase cho FCM")

    add_heading_styled(doc, "Các bước cài đặt Backend", level=2)
    add_code_block(doc,
        "# 1. Clone source\n"
        "git clone <repository-url>\n"
        "cd foot-delivery-android/backend\n\n"
        "# 2. Cài dependency\n"
        "pip install -r requirements.txt\n\n"
        "# 3. Cấu hình\n"
        "cp .env.example .env\n\n"
        "# 4. Chạy database\n"
        "docker-compose up -d postgres\n\n"
        "# 5. Khởi tạo database\n"
        "psql -U postgres -f main.sql\n\n"
        "# 6. Chạy backend\n"
        "uvicorn main:app --reload --host 0.0.0.0 --port 8000"
    )

    add_heading_styled(doc, "Các bước cài đặt Android App", level=2)
    add_code_block(doc,
        "# 1. Mở project trong Android Studio\n"
        "# File → Open → chọn thư mục foot-delivery-android/app\n\n"
        "# 2. Sync Gradle\n\n"
        "# 3. Cấu hình API URL (api/NetworkConfig.kt)\n"
        "# - Emulator: http://10.0.2.2:8000/\n"
        "# - Device: http://<your-ip>:8000/\n\n"
        "# 4. Chạy app: Run (Shift+F10)"
    )

    add_heading_styled(doc, "Lưu ý", level=2)
    add_bullet(doc, "Port: Backend chạy trên port 8000, emulator dùng 10.0.2.2")
    add_bullet(doc, "Biến môi trường: DATABASE_URL, JWT_SECRET_KEY, Firebase credentials")
    add_bullet(doc, "Firebase: Cần google-services.json từ Firebase Console")
    add_bullet(doc, "Lỗi thường gặp: Connection refused (backend chưa chạy), 401 (token hết hạn)")

    # ══════════════════════════════════════════════════════════
    # PHẦN 2: CODE
    # ══════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "PHẦN 2: CODE", level=0)

    # 1. Phần cá nhân
    add_heading_styled(doc, "1. Phần cá nhân thực hiện", level=1)

    add_table(doc,
        headers=["File", "Nội dung thực hiện", "Chức năng"],
        rows=[
            ["screens/home.kt", "Màn hình chính", "Hiển thị món ăn, tìm kiếm debounce 500ms"],
            ["screens/list_restaurant.kt", "Danh sách nhà hàng", "Tìm kiếm và hiển thị nhà hàng"],
            ["screens/restaurant_profile.kt", "Chi tiết nhà hàng", "Banner, thông tin, menu"],
            ["screens/food_detail.kt", "Chi tiết món ăn", "Đánh giá, số lượng, thêm giỏ"],
            ["screens/cart.kt", "Giỏ hàng", "Quản lý món, tính tổng"],
            ["screens/order.kt", "Xác nhận đơn", "Địa chỉ, voucher, thanh toán"],
            ["screens/OrderTracking.kt", "Theo dõi đơn", "Phân trang, phân loại pending/completed"],
            ["screens/OrderTrackingDetail.kt", "Chi tiết đơn + đánh giá", "Trạng thái, form đánh giá 1-5 sao"],
            ["screens/PointsDetail.kt", "Tích điểm", "Điểm theo đơn hàng"],
            ["screens/profile.kt", "Hồ sơ", "Hiển thị điểm tích lũy"],
            ["adapters/MenuItemAdapter.kt", "Adapter món ăn", "Hiển thị trên home"],
            ["adapters/RestaurantAdapter.kt", "Adapter nhà hàng", "Hiển thị danh sách"],
            ["api/ApiService.kt", "API definitions", "Tất cả data classes và endpoints"],
            ["api/RetrofitClient.kt", "HTTP Client", "Auth interceptor"],
        ],
        col_widths=[4, 3.5, 5]
    )

    # 2. Tối ưu source code
    add_heading_styled(doc, "2. Tối ưu source code", level=1)

    add_heading_styled(doc, "Đã tối ưu", level=2)
    add_table(doc,
        headers=["Vấn đề", "Giải pháp", "File"],
        rows=[
            ["Tìm kiếm gây nhiều request", "Debounce 500ms", "home.kt, list_restaurant.kt"],
            ["Load ảnh từ URL", "Picasso với placeholder + error", "MenuItemAdapter, RestaurantAdapter"],
            ["Auth token", "OkHttp Interceptor tự động đính kèm", "RetrofitClient.kt"],
            ["Giỏ hàng", "Static companion object (in-memory)", "cart.kt"],
        ],
        col_widths=[4, 5, 4]
    )

    add_heading_styled(doc, "Có thể bổ sung", level=2)
    add_table(doc,
        headers=["Vấn đề", "Đề xuất"],
        rows=[
            ["Không có pagination cho menu items", "Thêm infinite scroll"],
            ["Gọi getOrderDetail() N lần trong OrderTracking", "Batch API hoặc dùng response đã có"],
            ["LoyaltyPoint table không dùng", "Xóa hoặc tích hợp vào tính điểm"],
            ["Không cache API response", "Thêm OkHttp cache hoặc Room"],
            ["Cart mất khi đóng app", "Lưu vào SharedPreferences hoặc Room"],
        ],
        col_widths=[6, 6]
    )

    # 3. Kiểm tra comment
    add_heading_styled(doc, "3. Kiểm tra comment", level=1)

    add_heading_styled(doc, "Tình trạng hiện tại", level=2)
    add_table(doc,
        headers=["Loại", "Trạng thái", "Ghi chú"],
        rows=[
            ["Class/Activity", "Đã comment ✅", "KDoc mô tả vai trò, luồng hoạt động"],
            ["Hàm", "Chưa comment ❌", "Nên thêm KDoc cho hàm phức tạp"],
            ["API interface", "Đã comment ✅", "KDoc cho mỗi endpoint"],
            ["Data class", "Đã comment ✅", "KDoc + @property cho mọi trường"],
            ["Logic phức tạp", "Chưa comment ❌", "Debounce, pagination, payment flow"],
        ],
        col_widths=[3, 3, 6]
    )

    add_heading_styled(doc, "Vị trí nên thêm comment", level=2)
    add_table(doc,
        headers=["File", "Vị trí", "Loại comment"],
        rows=[
            ["home.kt", "searchMenuItems()", "KDoc: debounce logic"],
            ["cart.kt", "companion object { cartList }", "KDoc: giải thích static cart"],
            ["order.kt", "createOrderFromCart()", "KDoc: flow đặt hàng"],
            ["order.kt", "calculateDiscountAmount()", "KDoc: logic tính giảm giá"],
            ["OrderTracking.kt", "loadAllOrders()", "KDoc: phân trang"],
            ["OrderTrackingDetail.kt", "updateStarColors()", "KDoc: UI star rating"],
            ["OrderTrackingDetail.kt", "createReview()", "KDoc: validation + submit"],
            ["payment_methods.kt", "startVnPayPayment()", "KDoc: VNPay flow"],
            ["VNPayActivity.kt", "shouldOverrideUrlLoading()", "KDoc: intercept callback"],
        ],
        col_widths=[4, 4, 4]
    )

    # ══════════════════════════════════════════════════════════
    # PHỤ LỤC
    # ══════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "PHỤ LỤC", level=0)

    add_heading_styled(doc, "Trạng thái đơn hàng", level=1)

    add_table(doc,
        headers=["Trạng thái", "Tiếng Việt", "Mô tả"],
        rows=[
            ["pending", "Chờ xác nhận", "Đơn mới tạo"],
            ["paid", "Đã thanh toán", "Đã thanh toán VNPay"],
            ["confirmed", "Đã xác nhận", "Xác nhận từ nhà hàng"],
            ["delivering", "Đang giao", "Shipper đang giao"],
            ["completed", "Đã giao", "Giao hàng thành công"],
            ["cancelled", "Đã hủy", "Đơn bị hủy"],
        ],
        col_widths=[3, 3, 6]
    )

    add_heading_styled(doc, "Danh sách data classes", level=1)

    add_heading_styled(doc, "Đặt món", level=2)
    add_code_block(doc,
        "data class MenuItem(id, name, image_url, price, is_available, description,\n"
        "    restaurantid, categoryid, restaurant_name, reviews, avg_rating)\n"
        "data class Restaurant(id, name, image_url, address, rating, open_time,\n"
        "    close_time, phone_number, status, description, menu_items)\n"
        "data class CartItem(id, name, price, qty, imageUrl, isSelected)\n"
        "data class PreOrderItem(id, name, price, qty, imageUrl, deliveryTime, isSelected)"
    )

    add_heading_styled(doc, "Đặt hàng", level=2)
    add_code_block(doc,
        "data class OrderCreateRequest(status, preorderdate, preordertime, totalprice,\n"
        "    restaurantid, addressid, userid, order_items)\n"
        "data class OrderItemRequest(quantity, price, menuitemid)\n"
        "data class OrderResponse(id, status, createdat, preorderdate, preordertime,\n"
        "    totalprice, restaurantid, addressid, userid)\n"
        "data class OrderDetailResponse(id, status, createdat, preorderdate, preordertime,\n"
        "    totalprice, restaurantid, addressid, userid, restaurant_name, address_detail, order_items)"
    )

    add_heading_styled(doc, "Đánh giá", level=2)
    add_code_block(doc,
        "data class ReviewDetail(id, rating, comment, userid, user_name)\n"
        "data class ReviewCreateRequest(rating, comment, orderid, userid, menuitemid, restaurantid)\n"
        "data class ReviewResponse(id, rating, comment, orderid, menuitemid, restaurantid, userid)"
    )

    add_heading_styled(doc, "Người dùng", level=2)
    add_code_block(doc,
        "data class UserProfileSummary(user_id, user_name, points, delivered_orders, total_spent)\n"
        "data class Address(id, detail, phone, userid)\n"
        "data class LoginRequest(username, password)\n"
        "data class LoginResponse(access_token, token_type)"
    )

    # ── Lưu file ──────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(__file__), "bao-cao-ky-thuat.docx")
    doc.save(output_path)
    print(f"✅ Đã tạo file: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
